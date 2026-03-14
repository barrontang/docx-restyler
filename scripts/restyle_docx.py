#!/usr/bin/env python3
import argparse
import re
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


@dataclass
class ParagraphFormatSpec:
    style_name: str
    font_name: Optional[str] = None
    font_size_pt: Optional[float] = None
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    alignment: Optional[int] = None
    first_line_indent = None
    left_indent = None
    right_indent = None
    space_before = None
    space_after = None
    line_spacing = None
    char_spacing_twips: Optional[int] = None


@dataclass
class TemplateInfo:
    page_section = None
    title: ParagraphFormatSpec = None
    heading1: ParagraphFormatSpec = None
    heading2: ParagraphFormatSpec = None
    heading3: ParagraphFormatSpec = None
    heading4: ParagraphFormatSpec = None
    body: ParagraphFormatSpec = None
    quote: Optional[ParagraphFormatSpec] = None


def style_exists(doc: Document, style_name: str) -> bool:
    try:
        doc.styles[style_name]
        return True
    except KeyError:
        return False


def get_run_spacing(run) -> Optional[int]:
    rPr = run._element.rPr
    if rPr is None:
        return None
    spacing = rPr.find(qn('w:spacing'))
    if spacing is None:
        return None
    val = spacing.get(qn('w:val'))
    if val is None:
        return None
    try:
        return int(val)
    except Exception:
        return None


def set_run_spacing(run, twips: Optional[int]) -> None:
    if twips is None:
        return
    rPr = run._element.get_or_add_rPr()
    spacing = rPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        rPr.append(spacing)
    spacing.set(qn('w:val'), str(int(twips)))


def _pick_run_font(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            font = run.font
            name = font.name or getattr(font, 'east_asia', None)
            size = font.size.pt if font.size else None
            bold = run.bold
            italic = run.italic
            spacing = get_run_spacing(run)
            return name, size, bold, italic, spacing
    return None, None, None, None, None


def paragraph_spec_from_style(doc: Document, style_name: str) -> ParagraphFormatSpec:
    spec = ParagraphFormatSpec(style_name=style_name)
    try:
        style = doc.styles[style_name]
    except KeyError:
        return spec

    font = style.font
    spec.font_name = font.name
    spec.font_size_pt = font.size.pt if font.size else None
    spec.bold = font.bold
    spec.italic = font.italic
    pf = style.paragraph_format
    spec.alignment = pf.alignment
    spec.first_line_indent = pf.first_line_indent
    spec.left_indent = pf.left_indent
    spec.right_indent = pf.right_indent
    spec.space_before = pf.space_before
    spec.space_after = pf.space_after
    spec.line_spacing = pf.line_spacing
    return spec


def paragraph_spec_from_paragraph(paragraph, fallback_style_name: str) -> ParagraphFormatSpec:
    spec = ParagraphFormatSpec(style_name=fallback_style_name)
    name, size, bold, italic, spacing = _pick_run_font(paragraph)
    spec.font_name = name
    spec.font_size_pt = size
    spec.bold = bold
    spec.italic = italic
    spec.char_spacing_twips = spacing
    pf = paragraph.paragraph_format
    spec.alignment = pf.alignment
    spec.first_line_indent = pf.first_line_indent
    spec.left_indent = pf.left_indent
    spec.right_indent = pf.right_indent
    spec.space_before = pf.space_before
    spec.space_after = pf.space_after
    spec.line_spacing = pf.line_spacing
    return spec


def find_paragraph_by_pattern(doc: Document, pattern: str):
    rx = re.compile(pattern)
    for p in doc.paragraphs:
        text = (p.text or '').strip()
        if text and rx.match(text):
            return p
    return None


def first_meaningful_paragraph(doc: Document):
    for p in doc.paragraphs:
        if (p.text or '').strip():
            return p
    return None


def merge_spec(primary: ParagraphFormatSpec, fallback: ParagraphFormatSpec) -> ParagraphFormatSpec:
    out = deepcopy(fallback)
    for field in out.__dataclass_fields__:
        val = getattr(primary, field)
        if val is not None:
            setattr(out, field, val)
    return out


def extract_template_info(template_doc: Document) -> TemplateInfo:
    title_p = first_meaningful_paragraph(template_doc)
    h1_p = find_paragraph_by_pattern(template_doc, r'^[一二三四五六七八九十]+、')
    h2_p = find_paragraph_by_pattern(template_doc, r'^（[一二三四五六七八九十]+）')
    h3_p = find_paragraph_by_pattern(template_doc, r'^\d+[\.、]')
    h4_p = find_paragraph_by_pattern(template_doc, r'^（\d+）')
    body_p = None
    for p in template_doc.paragraphs:
        txt = (p.text or '').strip()
        if not txt:
            continue
        if txt in {
            title_p.text.strip() if title_p else '',
            h1_p.text.strip() if h1_p else '',
            h2_p.text.strip() if h2_p else '',
            h3_p.text.strip() if h3_p else '',
            h4_p.text.strip() if h4_p else '',
        }:
            continue
        if len(txt) > 20:
            body_p = p
            break

    normal_style = paragraph_spec_from_style(template_doc, 'Normal') if style_exists(template_doc, 'Normal') else ParagraphFormatSpec('Normal')
    body = merge_spec(paragraph_spec_from_paragraph(body_p, 'Normal'), normal_style) if body_p else normal_style
    title = merge_spec(paragraph_spec_from_paragraph(title_p, 'Title' if style_exists(template_doc, 'Title') else 'Normal'), paragraph_spec_from_style(template_doc, 'Title') if style_exists(template_doc, 'Title') else body) if title_p else body
    h1 = merge_spec(paragraph_spec_from_paragraph(h1_p, 'Heading 1' if style_exists(template_doc, 'Heading 1') else 'Normal'), paragraph_spec_from_style(template_doc, 'Heading 1') if style_exists(template_doc, 'Heading 1') else body) if h1_p else body
    h2 = merge_spec(paragraph_spec_from_paragraph(h2_p, 'Heading 2' if style_exists(template_doc, 'Heading 2') else 'Normal'), paragraph_spec_from_style(template_doc, 'Heading 2') if style_exists(template_doc, 'Heading 2') else h1) if h2_p else h1
    h3 = merge_spec(paragraph_spec_from_paragraph(h3_p, 'Heading 3' if style_exists(template_doc, 'Heading 3') else 'Normal'), paragraph_spec_from_style(template_doc, 'Heading 3') if style_exists(template_doc, 'Heading 3') else h2) if h3_p else h2
    h4 = merge_spec(paragraph_spec_from_paragraph(h4_p, 'Normal'), body) if h4_p else body
    quote = paragraph_spec_from_style(template_doc, 'Quote') if style_exists(template_doc, 'Quote') else None

    info = TemplateInfo()
    info.page_section = template_doc.sections[0] if template_doc.sections else None
    info.title = title
    info.heading1 = h1
    info.heading2 = h2
    info.heading3 = h3
    info.heading4 = h4
    info.body = body
    info.quote = quote
    return info


def copy_page_setup(src: Document, dst: Document) -> None:
    if not src.sections or not dst.sections:
        return
    src_sec = src.sections[0]
    dst_sec = dst.sections[0]
    dst_sec.page_width = src_sec.page_width
    dst_sec.page_height = src_sec.page_height
    dst_sec.left_margin = src_sec.left_margin
    dst_sec.right_margin = src_sec.right_margin
    dst_sec.top_margin = src_sec.top_margin
    dst_sec.bottom_margin = src_sec.bottom_margin
    dst_sec.header_distance = src_sec.header_distance
    dst_sec.footer_distance = src_sec.footer_distance


def classify_paragraph(p, idx: int, total: int, prev_text: str = '', next_text: str = '') -> str:
    text = (p.text or '').strip()
    if not text:
        return 'skip'
    style_name = getattr(getattr(p, 'style', None), 'name', '') or ''

    # 0. hard-rule official Chinese heading hierarchy: highest priority
    if re.match(r'^[一二三四五六七八九十百千万]+、', text):
        return 'heading1'
    if re.match(r'^（[一二三四五六七八九十百千万]+）', text):
        return 'heading2'
    if re.match(r'^\d+[\.、]', text):
        return 'heading3'
    if re.match(r'^（\d+）', text):
        return 'heading4'

    # 1. explicit style names if source file actually has useful styles
    if style_name == 'Title':
        return 'title'
    if style_name.startswith('Heading 1'):
        return 'heading1'
    if style_name.startswith('Heading 2'):
        return 'heading2'
    if style_name.startswith('Heading 3'):
        return 'heading3'
    if 'Quote' in style_name:
        return 'quote'

    # 2. title heuristic only for first line
    if idx == 0 and len(text) <= 80:
        return 'title'

    # 3. tail signature/date blocks
    if idx >= total - 2 and (re.search(r'\d{4}年\d{1,2}月\d{1,2}日', text) or len(text) <= 20):
        return 'tail'

    # 4. fallback only when not matching formal numbered headings
    if 4 <= len(text) <= 18 and len(next_text.strip()) >= 25 and not re.search(r'[。；！？]$', text):
        return 'heading4'

    return 'body'


def apply_paragraph_spec(para, spec: ParagraphFormatSpec, force_core: bool = True):
    pf = para.paragraph_format
    if spec.alignment is not None:
        para.alignment = spec.alignment
    if spec.first_line_indent is not None:
        pf.first_line_indent = spec.first_line_indent
    if spec.left_indent is not None:
        pf.left_indent = spec.left_indent
    if spec.right_indent is not None:
        pf.right_indent = spec.right_indent
    if spec.space_before is not None:
        pf.space_before = spec.space_before
    if spec.space_after is not None:
        pf.space_after = spec.space_after
    if spec.line_spacing is not None:
        pf.line_spacing = spec.line_spacing

    for run in para.runs:
        if force_core:
            if spec.font_name is not None:
                run.font.name = spec.font_name
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.rFonts
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                rFonts.set(qn('w:eastAsia'), spec.font_name)
                rFonts.set(qn('w:ascii'), spec.font_name)
                rFonts.set(qn('w:hAnsi'), spec.font_name)
            if spec.font_size_pt is not None:
                run.font.size = Pt(spec.font_size_pt)
            if spec.bold is not None:
                run.bold = spec.bold
            if spec.italic is not None:
                run.italic = spec.italic
            set_run_spacing(run, spec.char_spacing_twips)


def add_paragraph_with_spec(dst: Document, text: str, spec: ParagraphFormatSpec) -> None:
    para = dst.add_paragraph(text)
    apply_paragraph_spec(para, spec, force_core=True)


def copy_table(src_table, dst: Document, template_info: TemplateInfo) -> None:
    rows = len(src_table.rows)
    cols = len(src_table.columns)
    table = dst.add_table(rows=rows, cols=cols)
    try:
        table.style = src_table.style
    except Exception:
        pass
    for r_idx, row in enumerate(src_table.rows):
        for c_idx, cell in enumerate(row.cells):
            target = table.cell(r_idx, c_idx)
            target.text = cell.text
            for p in target.paragraphs:
                apply_paragraph_spec(p, template_info.body, force_core=True)


def rebuild_document(template_path: Path, source_path: Path, output_path: Path) -> None:
    template_doc = Document(str(template_path))
    source_doc = Document(str(source_path))
    output_doc = Document()

    copy_page_setup(template_doc, output_doc)
    info = extract_template_info(template_doc)

    if output_doc.paragraphs and not output_doc.paragraphs[0].text:
        p = output_doc.paragraphs[0]._element
        p.getparent().remove(p)

    src_paras = source_doc.paragraphs
    total = len(src_paras)
    for para_index, p in enumerate(src_paras):
        text = (p.text or '').strip()
        prev_text = (src_paras[para_index - 1].text or '').strip() if para_index > 0 else ''
        next_text = (src_paras[para_index + 1].text or '').strip() if para_index + 1 < total else ''
        kind = classify_paragraph(p, para_index, total, prev_text, next_text)
        if kind == 'skip':
            continue
        if kind == 'title':
            add_paragraph_with_spec(output_doc, text, info.title)
        elif kind == 'heading1':
            add_paragraph_with_spec(output_doc, text, info.heading1)
        elif kind == 'heading2':
            add_paragraph_with_spec(output_doc, text, info.heading2)
        elif kind == 'heading3':
            add_paragraph_with_spec(output_doc, text, info.heading3)
        elif kind == 'heading4':
            add_paragraph_with_spec(output_doc, text, info.heading4)
        elif kind == 'tail':
            add_paragraph_with_spec(output_doc, text, info.body)
            if output_doc.paragraphs:
                output_doc.paragraphs[-1].alignment = 2
        elif kind == 'quote' and info.quote:
            add_paragraph_with_spec(output_doc, text, info.quote)
        else:
            add_paragraph_with_spec(output_doc, text, info.body)

    for table in source_doc.tables:
        copy_table(table, output_doc, info)

    output_doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description='Restyle a source .docx using a template .docx')
    parser.add_argument('--template', required=True, help='Template .docx path')
    parser.add_argument('--source', required=True, help='Source .docx path')
    parser.add_argument('--output', required=True, help='Output .docx path')
    args = parser.parse_args()

    template_path = Path(args.template)
    source_path = Path(args.source)
    output_path = Path(args.output)

    if not template_path.exists():
        raise SystemExit(f'Template not found: {template_path}')
    if not source_path.exists():
        raise SystemExit(f'Source not found: {source_path}')

    rebuild_document(template_path, source_path, output_path)
    print(f'Wrote: {output_path}')


if __name__ == '__main__':
    main()
